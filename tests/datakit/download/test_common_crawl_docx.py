# Copyright The Marin Authors
# SPDX-License-Identifier: Apache-2.0

import io
import json
import zipfile
from dataclasses import dataclass

import pytest
from docx import Document
from marin.datakit.download.common_crawl_docx import (
    DOCX_MIME_TYPE,
    CommonCrawlDocxConfig,
    DoclingDocxExtractor,
    DocxRecordSelector,
    DocxSelectionReason,
    ExtractedDocument,
    InvalidDocxError,
    LanguageDetection,
    LinguaLanguageDetector,
    common_crawl_docx_steps,
    extracted_docx_record,
    process_fetched_docx,
    validate_docx,
)
from marin.datakit.download.common_crawl_plan import (
    CommonCrawlIndexKind,
    CommonCrawlSource,
    FetchedCommonCrawlRecord,
)
from marin.datakit.download.common_crawl_warc import CommonCrawlWarcRecord, content_digest, main_record_from_index_row

from experiments.datakit.common_crawl_docx_sample import sample_report_markdown, stratified_partition_slice

CRAWL_ID = "CC-MAIN-2026-30"
RECORD_ID = "<urn:uuid:019f8700-d21d-78d8-8eb1-99eaa22579da>"
URL = "https://example.com/report.docx"


def _docx_payload(*, document: bytes = b"<w:document/>") -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", b"<Types/>")
        archive.writestr("_rels/.rels", b"<Relationships/>")
        archive.writestr("word/document.xml", document)
    return output.getvalue()


def _real_docx_payload(text: str) -> bytes:
    output = io.BytesIO()
    document = Document()
    document.add_heading("Common Crawl DOCX", level=1)
    document.add_paragraph(text)
    document.save(output)
    return output.getvalue()


def _real_docx_with_table() -> bytes:
    output = io.BytesIO()
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Quarter"
    table.cell(0, 1).text = "Revenue"
    table.cell(1, 0).text = "Q1"
    table.cell(1, 1).text = "$42"
    document.save(output)
    return output.getvalue()


def _index_row(payload: bytes) -> dict[str, object]:
    return {
        "url": URL,
        "fetch_status": 200,
        "content_mime_type": DOCX_MIME_TYPE,
        "content_mime_detected": "application/zip",
        "content_digest": content_digest(payload),
        "content_truncated": None,
        "warc_filename": "crawl-data/test.warc.gz",
        "warc_record_offset": 42,
        "warc_record_length": 100,
        "warc_record_id": RECORD_ID,
    }


def _fetched(payload: bytes) -> FetchedCommonCrawlRecord:
    row = _index_row(payload)
    selection = DocxRecordSelector().select(row)
    assert selection is not None
    return FetchedCommonCrawlRecord(
        indexed_record=main_record_from_index_row(row, crawl_id=CRAWL_ID),
        selection=selection,
        observed_record=CommonCrawlWarcRecord(
            payload=payload,
            payload_digest=content_digest(payload),
            warc_record_id=RECORD_ID,
            target_url=URL,
            http_status=200,
            http_content_type=DOCX_MIME_TYPE,
            warc_date="2026-07-21T21:48:44Z",
            identified_payload_type="application/zip",
        ),
    )


def _config(**kwargs: object) -> CommonCrawlDocxConfig:
    return CommonCrawlDocxConfig(
        name="test-docx",
        sources=(
            CommonCrawlSource(
                crawl_id=CRAWL_ID,
                index_kind=CommonCrawlIndexKind.MAIN,
                paths_manifest_url="https://example.com/index.paths.gz",
            ),
        ),
        **kwargs,
    )


@dataclass(frozen=True)
class _Extractor:
    text: str
    version: str = "test-extractor-v1"

    def extract(self, payload: bytes) -> ExtractedDocument:
        assert payload.startswith(b"PK")
        return ExtractedDocument(text=self.text, word_count=len(self.text.split()), table_count=2, image_count=1)


@dataclass(frozen=True)
class _LanguageDetector:
    detection: LanguageDetection
    version: str = "test-detector-v1"

    def detect(self, text: str) -> LanguageDetection:
        assert text
        return self.detection


@pytest.mark.parametrize(
    ("row", "reason"),
    [
        ({"fetch_status": 200, "content_mime_type": DOCX_MIME_TYPE, "url": "https://example.com/a"}, "declared_mime"),
        ({"fetch_status": 200, "content_mime_type": "application/octet-stream", "url": URL}, "url_suffix"),
        (
            {"fetch_status": 200, "content_mime_detected": DOCX_MIME_TYPE, "url": "https://example.com/a"},
            "detected_mime",
        ),
    ],
)
def test_selector_preserves_selection_reason(row: dict[str, object], reason: str) -> None:
    selection = DocxRecordSelector().select(row)

    assert selection is not None
    assert selection.metadata["selection_reason"] == reason


def test_selector_rejects_failed_and_truncated_rows() -> None:
    selector = DocxRecordSelector()

    assert selector.select({"fetch_status": 404, "content_mime_type": DOCX_MIME_TYPE, "url": URL}) is None
    assert (
        selector.select(
            {"fetch_status": 200, "content_truncated": "length", "content_mime_type": DOCX_MIME_TYPE, "url": URL}
        )
        is None
    )


def test_validate_docx_accepts_required_office_members() -> None:
    validate_docx(_docx_payload(), maximum_entries=10, maximum_uncompressed_bytes=1024)


@pytest.mark.parametrize("payload", [b"not a zip", _docx_payload(document=b"x")[:-8]])
def test_validate_docx_rejects_malformed_payload(payload: bytes) -> None:
    with pytest.raises(InvalidDocxError):
        validate_docx(payload, maximum_entries=10, maximum_uncompressed_bytes=1024)


def test_validate_docx_enforces_entry_and_size_limits() -> None:
    payload = _docx_payload(document=b"x" * 100)

    with pytest.raises(InvalidDocxError, match="entries"):
        validate_docx(payload, maximum_entries=1, maximum_uncompressed_bytes=1024)
    with pytest.raises(InvalidDocxError, match="expands"):
        validate_docx(payload, maximum_entries=10, maximum_uncompressed_bytes=10)


def test_docling_extractor_reads_real_docx_fixture() -> None:
    extracted = DoclingDocxExtractor().extract(_real_docx_payload("A small document extracted by Docling."))

    assert "Common Crawl DOCX" in extracted.text
    assert "A small document extracted by Docling." in extracted.text


def test_docling_extractor_preserves_table_text_without_markdown_padding() -> None:
    extracted = DoclingDocxExtractor().extract(_real_docx_with_table())

    assert "Quarter | Revenue" in extracted.text
    assert "Q1 | $42" in extracted.text
    assert "---" not in extracted.text
    assert extracted.table_count == 1


def test_lingua_detector_identifies_multilingual_fixture() -> None:
    french = "Ceci est un document français avec suffisamment de mots pour identifier correctement sa langue. " * 3

    detection = LinguaLanguageDetector().detect(french)

    assert detection.language == "fr"
    assert detection.confidence > 0.5


def test_extracted_record_preserves_discovery_and_warc_provenance() -> None:
    output = extracted_docx_record(
        _fetched(_docx_payload()),
        extractor=_Extractor("Bonjour tout le monde."),
        language_detector=_LanguageDetector(LanguageDetection("fr", 0.98)),
        maximum_zip_entries=10,
        maximum_uncompressed_bytes=1024,
    )

    assert output["source_id"] == RECORD_ID
    assert output["crawl_id"] == CRAWL_ID
    assert output["warc_record_offset"] == 42
    assert output["index_status"] == 200
    assert output["index_content_type"] == DOCX_MIME_TYPE
    assert output["selection_reason"] == DocxSelectionReason.DECLARED_MIME.value
    assert output["language"] == "fr"
    assert output["table_count"] == 2


def test_document_local_extraction_failure_is_skipped() -> None:
    output = process_fetched_docx(
        _fetched(_docx_payload()),
        extractor=_Extractor("  \n"),
        language_detector=_LanguageDetector(LanguageDetection("unknown", 0.0)),
        maximum_zip_entries=10,
        maximum_uncompressed_bytes=1024,
    )

    assert output is None


def test_pipeline_places_shared_plan_between_discovery_and_extraction() -> None:
    discovery, plan, extraction, normalized = common_crawl_docx_steps(
        _config(index_batch_rows=7, max_workers=11),
        extractor=_Extractor("text", version="extractor-v7"),
        language_detector=_LanguageDetector(LanguageDetection("en", 1.0), version="detector-v3"),
    )

    assert discovery in plan.deps
    assert plan in extraction.deps
    assert extraction in normalized.deps
    assert "index_batch_rows" not in discovery.hash_attrs
    assert "max_workers" not in discovery.hash_attrs
    assert extraction.hash_attrs["extractor"] == "extractor-v7"
    assert extraction.hash_attrs["language_detector"] == "detector-v3"


def test_stratified_partition_slice_spans_manifest() -> None:
    partitions = tuple(f"part-{index}" for index in range(10))

    assert stratified_partition_slice(partitions, 4) == ("part-0", "part-3", "part-6", "part-9")


def test_sample_report_reads_selection_reason_from_shared_discovery_metadata() -> None:
    source = _config().sources[0]
    candidates = [
        {"selection_metadata": json.dumps({"selection_reason": DocxSelectionReason.DECLARED_MIME.value})},
        {"selection_metadata": json.dumps({"selection_reason": DocxSelectionReason.DECLARED_MIME.value})},
        {"selection_metadata": json.dumps({"selection_reason": DocxSelectionReason.URL_SUFFIX.value})},
    ]
    extracted = [
        {
            "selection_reason": DocxSelectionReason.DECLARED_MIME.value,
            "url": URL,
            "language": "en",
            "word_count": 4,
            "table_count": 1,
            "text": "A short extracted document.",
        }
    ]

    markdown, examples = sample_report_markdown(
        source=source,
        candidate_rows=candidates,
        extracted_rows=extracted,
        normalized_rows=extracted,
        extraction_counters={"common_crawl_docx/invalid_files": 2},
        examples_per_reason=1,
    )

    assert "| `declared_mime` | 2 | 1 | 50.0% |" in markdown
    assert "| `url_suffix` | 1 | 0 | 0.0% |" in markdown
    assert examples == [
        {
            "selection_reason": DocxSelectionReason.DECLARED_MIME.value,
            "url": URL,
            "language": "en",
            "word_count": 4,
            "table_count": 1,
            "excerpt": "A short extracted document.",
        }
    ]
